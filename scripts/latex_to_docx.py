import docx
from docx.shared import Pt, Inches
import re

def clean_latex(text):
    # Remove basic latex commands
    text = re.sub(r'\\textbf{(.*?)}', r'\1', text)
    text = re.sub(r'\\textit{(.*?)}', r'\1', text)
    text = re.sub(r'\\texttt{(.*?)}', r'\1', text)
    text = re.sub(r'\\mathcal{(.*?)}', r'\1', text)
    text = re.sub(r'\\sim', '~', text)
    text = re.sub(r'\\times', 'x', text)
    text = re.sub(r'\\%', '%', text)
    text = re.sub(r'\$', '', text)
    text = re.sub(r'\\\[', '', text)
    text = re.sub(r'\\\]', '', text)
    text = re.sub(r'\\begin{figure}.*?\\end{figure}', '[IMAGE PLACEHOLDER]', text, flags=re.DOTALL)
    text = re.sub(r'\\begin{table}.*?\\end{table}', '[TABLE PLACEHOLDER]', text, flags=re.DOTALL)
    text = re.sub(r'\\begin{.*?}', '', text)
    text = re.sub(r'\\end{.*?}', '', text)
    text = re.sub(r'\\label{.*?}', '', text)
    text = re.sub(r'\\caption{.*?}', '', text)
    text = re.sub(r'\n+', '\n\n', text)
    return text.strip()

doc = docx.Document()

# Add Title
doc.add_heading("Comprehensive Machine Learning and Geospatial Clustering Approach for Real-Time Dynamic Road Accident Risk Zone Detection and ADAS Alerting", 0)

doc.add_paragraph("Navadeep, Inderneel, Rithik Teja, Shanmukh\nGuide: Dr./Mr. Ravi Krishna\nAnurag University, Hyderabad")

# Read LaTeX
with open('docs/paper.tex', 'r', encoding='utf-8') as f:
    latex_content = f.read()

# Extract Abstract
abstract_match = re.search(r'\\begin{abstract}(.*?)\\end{abstract}', latex_content, re.DOTALL)
if abstract_match:
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(clean_latex(abstract_match.group(1)))

# Extract sections
sections = re.findall(r'\\section{(.*?)\}(.*?)(?=\\section|\\end{document})', latex_content, re.DOTALL)

for title, content in sections:
    doc.add_heading(title, level=1)
    
    # Check for subsections
    subsections = re.split(r'\\subsection{(.*?)\}', content)
    
    # The first item is text before any subsection
    if subsections[0].strip():
        doc.add_paragraph(clean_latex(subsections[0]))
        
    # Pairs of title, content
    for i in range(1, len(subsections), 2):
        sub_title = subsections[i]
        sub_content = subsections[i+1]
        
        doc.add_heading(sub_title, level=2)
        doc.add_paragraph(clean_latex(sub_content))

doc.save('docs/SafeRoute_AI_Research_Paper.docx')
print("Saved to docs/SafeRoute_AI_Research_Paper.docx")
